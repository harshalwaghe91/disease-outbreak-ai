from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs"
ASSET_DIR = OUT_DIR / "paper_assets"
OUTPUT = OUT_DIR / "IEEE_Research_Paper_Disease_Outbreak_AI.docx"

FONT = "Times New Roman"
BLACK = RGBColor(0, 0, 0)
GRAY = RGBColor(75, 75, 75)
LIGHT = (241, 244, 246)
MID = (83, 98, 108)
TEAL = (14, 116, 110)


def set_run_font(run, size=10, bold=False, italic=False, color=BLACK):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def set_cell_margins(cell, top=55, start=70, bottom=55, end=70):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_borders(table, color="808080", size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa):
    total = sum(widths_dxa)
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def configure_page(section):
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(0.625)
    section.right_margin = Inches(0.625)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.4)


def set_columns(section, count=2, space_twips=360):
    sect_pr = section._sectPr
    cols = sect_pr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sect_pr.append(cols)
    cols.set(qn("w:num"), str(count))
    cols.set(qn("w:space"), str(space_twips))
    cols.set(qn("w:equalWidth"), "1")


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(10)
    normal.font.color.rgb = BLACK
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(2)
    normal.paragraph_format.first_line_indent = Inches(0.14)

    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        style = styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.font.size = Pt(10)
        style.font.color.rgb = BLACK
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    styles["Heading 1"].font.bold = True
    styles["Heading 1"].font.italic = False
    styles["Heading 1"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    styles["Heading 1"].paragraph_format.space_before = Pt(6)
    styles["Heading 1"].paragraph_format.space_after = Pt(2)

    styles["Heading 2"].font.bold = False
    styles["Heading 2"].font.italic = True
    styles["Heading 2"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    styles["Heading 2"].paragraph_format.space_before = Pt(4)
    styles["Heading 2"].paragraph_format.space_after = Pt(1)

    styles["Heading 3"].font.bold = False
    styles["Heading 3"].font.italic = True
    styles["Heading 3"].paragraph_format.space_before = Pt(3)
    styles["Heading 3"].paragraph_format.space_after = Pt(1)


def add_body(doc, text, indent=True):
    p = doc.add_paragraph(style="Normal")
    if not indent:
        p.paragraph_format.first_line_indent = Inches(0)
    run = p.add_run(text)
    set_run_font(run, 10)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    for run in p.runs:
        set_run_font(run, 10, bold=(level == 1), italic=(level > 1))
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    r = p.add_run(text)
    set_run_font(r, 8)
    return p


def add_table(doc, caption, headers, rows, widths_dxa, font_size=8):
    add_caption(doc, caption)
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_geometry(table, widths_dxa)
    set_table_borders(table)
    header = table.rows[0]
    set_repeat_table_header(header)
    prevent_row_split(header)
    for idx, label in enumerate(headers):
        cell = header.cells[idx]
        set_cell_shading(cell, "E6E6E6")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.first_line_indent = Inches(0)
        r = p.add_run(label)
        set_run_font(r, font_size, bold=True)

    for row_data in rows:
        row = table.add_row()
        prevent_row_split(row)
        for idx, value in enumerate(row_data):
            cell = row.cells[idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.first_line_indent = Inches(0)
            r = p.add_run(str(value))
            set_run_font(r, font_size)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(0)
    after.paragraph_format.line_spacing = 1
    return table


def font_for_diagram(size, bold=False):
    names = ["timesbd.ttf", "times.ttf"] if bold else ["times.ttf"]
    for name in names:
        path = Path("C:/Windows/Fonts") / name
        if path.exists():
            return ImageFont.truetype(str(path), size)
    return ImageFont.load_default()


def rounded_box(draw, xy, title, subtitle, fill, outline=TEAL):
    draw.rounded_rectangle(xy, radius=18, fill=fill, outline=outline, width=4)
    x1, y1, x2, y2 = xy
    title_font = font_for_diagram(30, bold=True)
    sub_font = font_for_diagram(22)
    title_box = draw.textbbox((0, 0), title, font=title_font)
    draw.text(((x1 + x2 - (title_box[2] - title_box[0])) / 2, y1 + 22), title, font=title_font, fill=(20, 30, 35))
    subtitle_box = draw.multiline_textbbox((0, 0), subtitle, font=sub_font, spacing=4, align="center")
    draw.multiline_text(
        ((x1 + x2 - (subtitle_box[2] - subtitle_box[0])) / 2, y1 + 66),
        subtitle,
        font=sub_font,
        fill=(45, 55, 60),
        spacing=4,
        align="center",
    )


def arrow(draw, start, end):
    draw.line([start, end], fill=MID, width=6)
    ex, ey = end
    sx, sy = start
    if abs(ex - sx) >= abs(ey - sy):
        direction = 1 if ex > sx else -1
        points = [(ex, ey), (ex - 18 * direction, ey - 12), (ex - 18 * direction, ey + 12)]
    else:
        direction = 1 if ey > sy else -1
        points = [(ex, ey), (ex - 12, ey - 18 * direction), (ex + 12, ey - 18 * direction)]
    draw.polygon(points, fill=MID)


def make_architecture(path):
    image = Image.new("RGB", (1200, 850), "white")
    draw = ImageDraw.Draw(image)
    rounded_box(draw, (70, 80, 440, 245), "React Client", "Forms, dashboard,\nreports and chatbot", LIGHT)
    rounded_box(draw, (760, 80, 1130, 245), "FastAPI Service", "Validation, routing,\nCORS and responses", LIGHT)
    rounded_box(draw, (70, 550, 440, 715), "SQLite Store", "Prediction history\nand dashboard records", LIGHT)
    rounded_box(draw, (760, 550, 1130, 715), "ML Services", "Random Forest models\nand clinical rules", LIGHT)
    rounded_box(draw, (420, 315, 780, 475), "REST API", "POST /predict, POST /chat\nGET analytics and reports", (225, 244, 242))
    arrow(draw, (440, 162), (760, 162))
    arrow(draw, (945, 245), (945, 550))
    arrow(draw, (760, 635), (440, 635))
    arrow(draw, (255, 550), (255, 245))
    arrow(draw, (440, 395), (420, 395))
    arrow(draw, (780, 395), (760, 395))
    image.save(path, dpi=(300, 300))


def make_pipeline(path):
    image = Image.new("RGB", (1200, 480), "white")
    draw = ImageDraw.Draw(image)
    boxes = [
        (25, 120, 240, 340, "Input", "13 fields"),
        (265, 120, 480, 340, "Transform", "Scale, encode,\nTF-IDF"),
        (505, 120, 720, 340, "Classify", "Two 220-tree\nforests"),
        (745, 120, 960, 340, "Adjust", "Clinical\nseverity rules"),
        (985, 120, 1175, 340, "Output", "Risk, class,\nprobability"),
    ]
    for x1, y1, x2, y2, title, subtitle in boxes:
        rounded_box(draw, (x1, y1, x2, y2), title, subtitle, LIGHT)
    for idx in range(len(boxes) - 1):
        arrow(draw, (boxes[idx][2], 230), (boxes[idx + 1][0], 230))
    image.save(path, dpi=(300, 300))


def add_figure(doc, image_path, caption, width=3.18):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(0)
    p.add_run().add_picture(str(image_path), width=Inches(width))
    add_caption(doc, caption)


def add_reference(doc, number, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.first_line_indent = Inches(-0.18)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    r = p.add_run(f"[{number}] {text}")
    set_run_font(r, 8)


def add_title_block(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("AI-Based Disease Outbreak Prediction Using Laboratory Test Data\nwith a General Consultation Chatbot")
    set_run_font(r, 24, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("Harshal Gulabrao Waghe")
    set_run_font(r, 11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(7)
    r = p.add_run(
        "Electronics and telecommunication department\n"
        "Suryodaya college of engineering and technology, Nagpur\n"
        "Nagpur, Maharashtra\n"
        "harshalwaghe91@gmail.com"
    )
    set_run_font(r, 10)

    abstract = (
        "Early recognition of abnormal laboratory patterns can support timely public-health investigation, "
        "but many academic prototypes either expose a machine-learning model without an accessible interface "
        "or provide dashboards without an end-to-end prediction workflow. This paper presents OutbreakAI, a "
        "full-stack decision-support prototype that combines laboratory measurements, demographics, symptoms, "
        "machine learning, analytics, report storage, and a safety-oriented consultation chatbot. A synthetic "
        "dataset of 5,000 records was generated across five disease categories and three risk levels. Numeric "
        "features are standardized, categorical variables are one-hot encoded, symptom text is represented by "
        "term frequency-inverse document frequency features, and two Random Forest classifiers containing 220 "
        "trees each predict disease category and risk level. A transparent clinical-severity layer combines model "
        "confidence with abnormal temperature, oxygen saturation, inflammatory markers, white blood cell count, "
        "platelets, and blood sugar. On a held-out set of 1,000 synthetic records, the disease classifier obtained "
        "1.000 accuracy and the risk classifier obtained 0.984 accuracy. These results demonstrate software and "
        "pipeline feasibility, not clinical validity; prospective evaluation on representative, expert-labeled "
        "data is required before any healthcare use."
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    label = p.add_run("Abstract—")
    set_run_font(label, 9, bold=True, italic=True)
    r = p.add_run(abstract)
    set_run_font(r, 9)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0)
    p.paragraph_format.space_after = Pt(6)
    label = p.add_run("Keywords—")
    set_run_font(label, 9, bold=True, italic=True)
    r = p.add_run(
        "disease outbreak prediction, laboratory test data, Random Forest, public-health surveillance, "
        "FastAPI, healthcare chatbot, synthetic data."
    )
    set_run_font(r, 9)


def build():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    architecture = ASSET_DIR / "system_architecture.png"
    pipeline = ASSET_DIR / "prediction_pipeline.png"
    make_architecture(architecture)
    make_pipeline(pipeline)

    doc = Document()
    configure_styles(doc)
    configure_page(doc.sections[0])
    set_columns(doc.sections[0], 1)
    add_title_block(doc)

    main_section = doc.add_section(WD_SECTION.CONTINUOUS)
    configure_page(main_section)
    set_columns(main_section, 2, 360)

    add_heading(doc, "I. INTRODUCTION", 1)
    add_body(
        doc,
        "Disease surveillance depends on timely collection, analysis, interpretation, and communication of "
        "health information. Laboratory observations are particularly valuable because changes in white blood "
        "cell count, inflammatory markers, platelet count, oxygen saturation, and temperature may provide early "
        "signals of infectious or respiratory disease. The World Health Organization (WHO) describes early-warning "
        "systems as mechanisms that collect both immediate alerts and periodic surveillance data to support "
        "verification and response [1], [2]. However, laboratory values alone do not establish an outbreak or an "
        "individual diagnosis; interpretation requires epidemiological context, validated thresholds, and expert review."
    )
    add_body(
        doc,
        "This work develops a deployable academic prototype that converts laboratory and symptom inputs into a "
        "risk category, likely disease group, outbreak probability, and preventive recommendation. The system also "
        "maintains prediction history, visualizes aggregate trends, and provides a rule-based chatbot that supplies "
        "general guidance while escalating emergency terms. The contribution is an integrated, reproducible stack "
        "rather than a claim of clinical deployment."
    )
    add_body(
        doc,
        "The principal contributions are: (1) a multimodal preprocessing pipeline for numeric, categorical, and "
        "free-text inputs; (2) separate Random Forest models for disease category and risk classification; (3) a "
        "transparent rule layer that adjusts categories and probability using clinically interpretable abnormalities; "
        "(4) a full-stack interface with dashboards, reports, and chatbot safety controls; and (5) a deployment design "
        "using free-tier web services."
    )

    add_heading(doc, "II. RELATED WORK", 1)
    add_body(
        doc,
        "Random Forest is an ensemble method that combines randomized decision trees and aggregates their predictions. "
        "Breiman showed that the method can provide strong generalization while remaining comparatively robust to "
        "noise [3]. It is suitable for tabular laboratory data because it models nonlinear interactions, tolerates "
        "mixed feature scales after preprocessing, and provides class probabilities. The implementation uses "
        "scikit-learn's RandomForestClassifier [4], supported by pipelines and column-wise transformations [5]."
    )
    add_body(
        doc,
        "Prior public-health systems emphasize that alerts should initiate investigation rather than act as final "
        "diagnoses [1]. OutbreakAI follows this principle by presenting outputs as decision-support estimates, "
        "displaying a medical disclaimer, and separating emergency chatbot messages from routine advice. The system "
        "also preserves a clear distinction between patient-level classification and population-level outbreak "
        "surveillance: the dashboard summarizes submitted prototype records but does not estimate real incidence."
    )

    add_heading(doc, "III. SYSTEM DESIGN AND METHODOLOGY", 1)
    add_heading(doc, "A. System Architecture", 2)
    add_body(
        doc,
        "The application follows a client-server architecture. A React single-page application sends JSON requests "
        "through Axios to a FastAPI backend. The backend validates inputs, loads a serialized model bundle with joblib, "
        "stores prediction records in SQLite, and returns JSON responses. Cross-origin resource sharing is configured "
        "for local and deployed frontend origins [6]. Figure 1 summarizes the primary components."
    )
    add_figure(doc, architecture, "Fig. 1. OutbreakAI full-stack system architecture.")

    add_heading(doc, "B. Dataset Generation", 2)
    add_body(
        doc,
        "Because no de-identified institutional dataset was available, a reproducible generator produced 5,000 "
        "synthetic records using a fixed random seed of 42. Each record contains age, gender, location, symptom text, "
        "and ten laboratory or physiological measurements. Records are assigned to Viral Infection, Bacterial "
        "Infection, Respiratory Disease, Dengue-like Illness, or Flu-like Illness. Class-specific perturbations create "
        "high WBC and CRP for bacterial cases, lower oxygen saturation for respiratory cases, lower platelet counts "
        "for dengue-like cases, and increased temperature for dengue-like and flu-like cases."
    )
    add_table(
        doc,
        "TABLE I\nMODEL INPUT GROUPS",
        ["Group", "Variables"],
        [
            ("Demographic", "Age, gender, location"),
            ("Symptoms", "Free-text symptom description"),
            ("Hematology", "WBC, RBC, platelets, hemoglobin"),
            ("Inflammation", "CRP, ESR"),
            ("Vitals", "Temperature, oxygen saturation"),
            ("Metabolic", "Blood sugar"),
        ],
        [1500, 3200],
    )
    add_body(
        doc,
        "Risk labels are derived from five severity indicators: temperature above 39 °C, oxygen saturation below 93%, "
        "CRP above 55 mg/L, platelet count below 110×10³/µL, and WBC count above 15×10³/µL. Zero indicators produce "
        "Low risk, one or two produce Medium risk, and three or more produce High risk. This deterministic labeling "
        "supports pipeline testing but also makes the evaluation easier than a real clinical task."
    )

    add_heading(doc, "C. Preprocessing and Model Training", 2)
    add_body(
        doc,
        "The dataset is divided into 4,000 training and 1,000 testing records using an 80:20 stratified split. Numeric "
        "features are standardized, gender and location are one-hot encoded with unknown-category handling, and "
        "symptoms are transformed using TF-IDF with unigrams and bigrams and a maximum of 80 terms. Two independent "
        "Random Forest classifiers are trained with 220 trees, balanced class weights, parallel execution, and a fixed "
        "random state. The fitted pipelines, metrics, training time, and dataset size are saved in outbreak_model.pkl."
    )
    add_figure(doc, pipeline, "Fig. 2. Prediction preprocessing and inference pipeline.")

    add_heading(doc, "D. Hybrid Clinical Adjustment", 2)
    add_body(
        doc,
        "A clinical-severity score S complements the model output. Each input abnormality is normalized to [0,1] and "
        "combined as S = 0.20T + 0.25O + 0.15C + 0.10E + 0.12W + 0.13P + 0.05G, where T is temperature elevation, "
        "O is oxygen deficit, C is CRP elevation, E is ESR elevation, W is WBC deviation, P is platelet reduction, "
        "and G is blood-sugar abnormality. The displayed probability combines an 0.08 baseline, 0.78S, 0.08 model "
        "confidence, and an adjustment of 0.00, 0.10, or 0.20 for the model's Low, Medium, or High risk prediction."
    )
    add_body(
        doc,
        "Transparent category rules prioritize recognizable patterns: fever with thrombocytopenia suggests a "
        "dengue-like category; low oxygen or breathing symptoms suggest respiratory disease; WBC above 12×10³/µL "
        "or CRP above 35 mg/L suggests bacterial infection; and fever with cough, sore throat, or body ache suggests "
        "flu-like illness. These rules are educational heuristics and are not diagnostic criteria."
    )

    add_heading(doc, "E. Consultation Chatbot", 2)
    add_body(
        doc,
        "The chatbot is intentionally rule based so that its behavior remains deterministic and auditable. It provides "
        "general advice for fever, respiratory symptoms, headache, gastrointestinal symptoms, and dengue-like symptoms. "
        "Messages containing chest pain, breathing difficulty, unconsciousness, seizure, facial drooping, slurred "
        "speech, or sudden weakness trigger an emergency response directing the user to local emergency services. "
        "Every response includes a disclaimer and avoids final diagnosis."
    )

    add_heading(doc, "IV. IMPLEMENTATION", 1)
    add_heading(doc, "A. Backend and Database", 2)
    add_body(
        doc,
        "FastAPI exposes POST /predict and POST /chat together with GET /dashboard, GET /reports, and GET /analytics. "
        "Pydantic schemas validate request fields and response types. The model bundle is loaded once through a cached "
        "service function, reducing repeated disk access. SQLite provides a compact, serverless, single-file database "
        "appropriate for an academic prototype [7]. Prediction history supplies the reports table and aggregate "
        "dashboard statistics."
    )
    add_table(
        doc,
        "TABLE II\nREST API ENDPOINTS",
        ["Method", "Endpoint", "Purpose"],
        [
            ("POST", "/predict", "Generate and store prediction"),
            ("POST", "/chat", "Return general consultation"),
            ("GET", "/dashboard", "Summary metrics"),
            ("GET", "/reports", "Prediction history"),
            ("GET", "/analytics", "Chart-ready aggregates"),
        ],
        [800, 1350, 2550],
    )

    add_heading(doc, "B. Frontend and Deployment", 2)
    add_body(
        doc,
        "The React interface contains a landing page, prediction form, analytics dashboard, consultation chatbot, "
        "reports page, and administrative overview. Tailwind CSS supplies responsive styling, Recharts renders "
        "interactive charts, React Router manages navigation, and Axios handles API calls. The production frontend is "
        "deployed on Vercel and the backend on Render. The API timeout is set to 70 seconds to accommodate possible "
        "free-tier cold starts, while the interface reports request errors to the user instead of failing silently."
    )
    add_body(
        doc,
        "CORS allows the configured production domain and local development origins. Environment variables separate "
        "the backend URL from source code. The public deployment and repository demonstrate reproducibility, but "
        "free-tier instance sleep can increase first-request latency and local SQLite persistence may not be durable "
        "across every hosting lifecycle."
    )

    add_heading(doc, "V. EXPERIMENTAL SETUP", 1)
    add_body(
        doc,
        "Experiments use the deterministic synthetic dataset and fixed train-test partition described above. Weighted "
        "precision, recall, and F1-score are reported to account for class distribution. Accuracy measures the fraction "
        "of correct predictions. The disease and risk models are evaluated independently on the same held-out 1,000 "
        "records. No test examples are used for fitting."
    )
    add_body(
        doc,
        "Functional tests additionally submit low-risk routine values, bacterial-like high-WBC/high-CRP values, "
        "respiratory-like low-oxygen values, and dengue-like low-platelet values through the deployed prediction page. "
        "This verifies request validation, inference, response rendering, and report storage across the end-to-end stack."
    )

    add_heading(doc, "VI. RESULTS AND DISCUSSION", 1)
    add_table(
        doc,
        "TABLE III\nHELD-OUT MODEL PERFORMANCE",
        ["Target", "Accuracy", "Precision", "Recall", "F1"],
        [
            ("Disease", "1.0000", "1.0000", "1.0000", "1.0000"),
            ("Risk", "0.9840", "0.9843", "0.9840", "0.9796"),
        ],
        [1100, 900, 900, 900, 900],
    )
    add_body(
        doc,
        "Table III shows perfect disease-category classification and 98.4% risk accuracy on the held-out synthetic set. "
        "The disease result is explained by the data generator: each category receives a fixed symptom phrase and "
        "category-specific numeric shifts, creating highly separable classes. Consequently, 1.000 accuracy should not "
        "be interpreted as evidence of equivalent performance on noisy, incomplete, or overlapping clinical cases."
    )
    add_body(
        doc,
        "The risk model's lower F1-score indicates that minority or boundary cases are more difficult even under "
        "synthetic labeling. The hybrid layer improves visible responsiveness to clinically recognizable patterns. "
        "For example, a demonstration record with WBC 17×10³/µL, CRP 92 mg/L, ESR 58 mm/hr, temperature 39.4 °C, "
        "and oxygen saturation 96% returned Bacterial Infection, Medium risk, and a 54.72% outbreak probability. "
        "This case confirms that the corrected category logic no longer collapses diverse inputs into one class."
    )
    add_table(
        doc,
        "TABLE IV\nEXAMPLE END-TO-END OUTPUT",
        ["Field", "Value"],
        [
            ("Key inputs", "WBC 17; CRP 92; ESR 58; Temp. 39.4"),
            ("Predicted category", "Bacterial Infection"),
            ("Risk level", "Medium"),
            ("Probability", "54.72%"),
        ],
        [1800, 2900],
    )
    add_body(
        doc,
        "The dashboard, analytics charts, report history, and chatbot complete the intended functional requirements. "
        "Because the output probability is a designed blend rather than a statistically calibrated outbreak prevalence, "
        "it should be described as a prototype score. Calibration with epidemiological labels and regional case counts "
        "is necessary before the term probability can carry operational public-health meaning."
    )

    add_heading(doc, "VII. LIMITATIONS AND ETHICAL CONSIDERATIONS", 1)
    add_body(
        doc,
        "The primary limitation is synthetic data. The generator cannot reproduce measurement error, missing values, "
        "comorbidities, medication effects, demographic imbalance, disease co-occurrence, laboratory reference-range "
        "variation, or temporal and geographic clustering. The five broad categories are not formal diagnoses, and a "
        "single patient record cannot establish an outbreak. The system has not undergone external validation, "
        "prospective testing, clinician review, probability calibration, fairness analysis, or regulatory assessment."
    )
    add_body(
        doc,
        "Ethical deployment would require de-identification, explicit consent or lawful public-health authority, access "
        "control, encryption, retention policies, audit logs, subgroup performance reporting, and clinical governance. "
        "Human review must remain authoritative. Emergency warnings should be localized to the user's jurisdiction, "
        "and the chatbot must never replace professional evaluation. These safeguards are reflected in the prototype's "
        "disclaimers but are not fully implemented as production controls."
    )

    add_heading(doc, "VIII. CONCLUSION AND FUTURE WORK", 1)
    add_body(
        doc,
        "This paper presented OutbreakAI, a complete web-based disease-outbreak decision-support prototype integrating "
        "laboratory data, symptom text, Random Forest models, transparent severity rules, analytics, reports, and a "
        "general consultation chatbot. The implementation demonstrates that heterogeneous inputs can be processed and "
        "served through a responsive, deployable architecture. Held-out synthetic results were strong, but their "
        "interpretation is deliberately limited to pipeline feasibility."
    )
    add_body(
        doc,
        "Future work should replace synthetic labels with de-identified, expert-annotated multicenter data; introduce "
        "missing-data handling and uncertainty intervals; calibrate probabilities; evaluate temporal and geographic "
        "forecasting; compare Random Forest with gradient boosting and sequence models; add explainability using SHAP; "
        "and conduct fairness, usability, security, and prospective clinical studies. A population-level forecasting "
        "module should incorporate case counts and time windows rather than inferring outbreaks from isolated records."
    )

    add_heading(doc, "REFERENCES", 1)
    references = [
        "World Health Organization, WHO Recommended Surveillance Standards, 2nd ed. Geneva, Switzerland: WHO, 1999.",
        "World Health Organization Regional Office for the Eastern Mediterranean, “Early Warning, Alert and Response Network (EWARN),” [Online]. Available: https://www.emro.who.int/health-topics/ewarn/. Accessed: Jun. 27, 2026.",
        "L. Breiman, “Random forests,” Machine Learning, vol. 45, no. 1, pp. 5–32, 2001, doi: 10.1023/A:1010933404324.",
        "Scikit-learn Developers, “RandomForestClassifier,” scikit-learn documentation, [Online]. Available: https://scikit-learn.org/stable/modules/generated/sklearn.ensemble.RandomForestClassifier.html. Accessed: Jun. 27, 2026.",
        "F. Pedregosa et al., “Scikit-learn: Machine learning in Python,” Journal of Machine Learning Research, vol. 12, pp. 2825–2830, 2011.",
        "FastAPI, “CORS (Cross-Origin Resource Sharing),” [Online]. Available: https://fastapi.tiangolo.com/tutorial/cors/. Accessed: Jun. 27, 2026.",
        "SQLite Consortium, “About SQLite,” [Online]. Available: https://www.sqlite.org/about.html. Accessed: Jun. 27, 2026.",
        "React Team, “React Quick Start,” [Online]. Available: https://react.dev/learn. Accessed: Jun. 27, 2026.",
        "W. McKinney, “Data structures for statistical computing in Python,” in Proc. 9th Python in Science Conf., Austin, TX, USA, 2010, pp. 56–61.",
    ]
    for idx, reference in enumerate(references, start=1):
        add_reference(doc, idx, reference)

    # Keep all document defaults and language metadata aligned with the required typeface.
    styles = doc.styles
    for style in styles:
        if style.type == 1 and style._element.rPr is not None:
            style.font.name = FONT
            style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
            style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)

    doc.core_properties.title = "AI-Based Disease Outbreak Prediction Using Laboratory Test Data"
    doc.core_properties.subject = "IEEE-style final-year project research paper"
    doc.core_properties.author = "Harshal Gulabrao Waghe"
    doc.core_properties.keywords = "outbreak prediction, Random Forest, laboratory data, FastAPI, chatbot"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
